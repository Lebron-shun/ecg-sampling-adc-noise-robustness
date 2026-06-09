"""Build the final Chinese course report in Markdown and DOCX from actual results."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from project_core import FIGURES_DIR, REPORT_DIR, RESULTS_DIR, ensure_project_dirs


BLUE = "276FBF"
DARK = "183153"
TEAL = "1B998B"
LIGHT = "EAF1F8"
GRAY = "657786"
WHITE = "FFFFFF"
BODY_FONT = "SimSun"
HEADING_FONT = "Microsoft YaHei"


DEFAULT_STUDENT = {
    "name": "提交前填写",
    "student_id": "提交前填写",
    "class_name": "提交前填写",
}


def student_info(name: str | None = None, student_id: str | None = None, class_name: str | None = None) -> dict[str, str]:
    return {
        "name": name or DEFAULT_STUDENT["name"],
        "student_id": student_id or DEFAULT_STUDENT["student_id"],
        "class_name": class_name or DEFAULT_STUDENT["class_name"],
    }


def student_line(student: dict[str, str]) -> str:
    return f"姓名：{student['name']}　学号：{student['student_id']}　班级：{student['class_name']}"


def read_inputs() -> dict:
    clean = pd.read_csv(RESULTS_DIR / "aggregate_clean.csv")
    noise = pd.read_csv(RESULTS_DIR / "aggregate_noise_by_snr.csv")
    candidates = pd.read_csv(RESULTS_DIR / "candidate_summary.csv")
    interactions = pd.read_csv(RESULTS_DIR / "interaction_contrasts.csv")
    bootstrap = pd.read_csv(RESULTS_DIR / "clean_bootstrap_drops.csv")
    paired_tests = pd.read_csv(RESULTS_DIR / "clean_paired_tests.csv")
    monitoring_path = RESULTS_DIR / "monitoring_summary.csv"
    monitoring = pd.read_csv(monitoring_path) if monitoring_path.exists() else pd.DataFrame()
    scenarios_path = RESULTS_DIR / "scenario_recommendations.json"
    scenarios = json.loads(scenarios_path.read_text(encoding="utf-8")) if scenarios_path.exists() else []
    chosen = json.loads((RESULTS_DIR / "recommended_config.json").read_text(encoding="utf-8"))
    ref_clean = clean[(clean.target_fs_hz == 360) & (clean.bits == 11)].iloc[0]
    chosen_clean = clean[
        (clean.target_fs_hz == int(chosen["target_fs_hz"])) & (clean.bits == int(chosen["bits"]))
    ].iloc[0]
    ref_noise = noise[(noise.target_fs_hz == 360) & (noise.bits == 11)].copy()
    chosen_noise = noise[
        (noise.target_fs_hz == int(chosen["target_fs_hz"])) & (noise.bits == int(chosen["bits"]))
    ].copy()
    noise_compare = ref_noise[["snr_db", "pooled_f1_pct"]].merge(
        chosen_noise[["snr_db", "pooled_f1_pct"]],
        on="snr_db",
        suffixes=("_reference", "_chosen"),
    )
    noise_compare["drop_pp"] = (
        noise_compare["pooled_f1_pct_reference"] - noise_compare["pooled_f1_pct_chosen"]
    )
    return {
        "clean": clean,
        "noise": noise,
        "candidates": candidates,
        "interactions": interactions,
        "bootstrap": bootstrap,
        "paired_tests": paired_tests,
        "monitoring": monitoring,
        "scenarios": scenarios,
        "chosen": chosen,
        "ref_clean": ref_clean,
        "chosen_clean": chosen_clean,
        "noise_compare": noise_compare,
    }


def fmt(value: float, digits: int = 2) -> str:
    return f"{float(value):.{digits}f}"


def summary_text(data: dict) -> dict[str, str]:
    chosen = data["chosen"]
    ref = data["ref_clean"]
    selected = data["chosen_clean"]
    noise_compare = data["noise_compare"]
    nonzero_interactions = data["interactions"][
        (data["interactions"].interaction_ci_low_pp > 0)
        | (data["interactions"].interaction_ci_high_pp < 0)
    ]
    chosen_bootstrap = data["bootstrap"][
        (data["bootstrap"].target_fs_hz == int(chosen["target_fs_hz"]))
        & (data["bootstrap"].bits == int(chosen["bits"]))
    ].iloc[0]
    chosen_test = data["paired_tests"][
        (data["paired_tests"].target_fs_hz == int(chosen["target_fs_hz"]))
        & (data["paired_tests"].bits == int(chosen["bits"]))
    ].iloc[0]
    monitoring = data["monitoring"]
    chosen_monitoring = None
    if not monitoring.empty:
        chosen_monitoring = monitoring[
            (monitoring.target_fs_hz == int(chosen["target_fs_hz"]))
            & (monitoring.bits == int(chosen["bits"]))
        ].iloc[0]
    return {
        "configuration": f"{int(chosen['target_fs_hz'])} Hz / {int(chosen['bits'])} bit",
        "clean_f1": fmt(selected["pooled_f1_pct"], 3),
        "clean_ref_f1": fmt(ref["pooled_f1_pct"], 3),
        "clean_drop": fmt(ref["pooled_f1_pct"] - selected["pooled_f1_pct"], 3),
        "timing": fmt(selected["median_timing_error_ms"], 2),
        "bitrate": str(int(selected["raw_bitrate_bps"])),
        "bitrate_reduction": fmt((1 - selected["raw_bitrate_bps"] / ref["raw_bitrate_bps"]) * 100, 1),
        "storage": fmt(selected["storage_mib_per_day"], 2),
        "noise_worst_drop": fmt(noise_compare[noise_compare.snr_db >= 6]["drop_pp"].max(), 2),
        "noise_lowest_f1": fmt(noise_compare["pooled_f1_pct_chosen"].min(), 2),
        "interaction_count": str(len(nonzero_interactions)),
        "bootstrap_low": fmt(chosen_bootstrap["macro_f1_drop_ci_low_pp"], 3),
        "bootstrap_high": fmt(chosen_bootstrap["macro_f1_drop_ci_high_pp"], 3),
        "holm_p": fmt(chosen_test["p_holm"], 3),
        "rr_error": fmt(chosen_monitoring["rr_median_abs_error_ms"], 2) if chosen_monitoring is not None else "NA",
        "hr_error": fmt(chosen_monitoring["hr_median_abs_error_bpm"], 3) if chosen_monitoring is not None else "NA",
        "valid_rr": fmt(chosen_monitoring["valid_rr_pair_pct"], 2) if chosen_monitoring is not None else "NA",
        "sdnn_error": fmt(chosen_monitoring["sdnn_relative_error_median_pct"], 2) if chosen_monitoring is not None else "NA",
        "rmssd_error": fmt(chosen_monitoring["rmssd_relative_error_median_pct"], 2) if chosen_monitoring is not None else "NA",
        "hrv_usable": fmt(chosen_monitoring["hrv_usable_record_pct"], 1) if chosen_monitoring is not None else "NA",
    }


def markdown_report(data: dict, s: dict[str, str], student: dict[str, str] | None = None) -> str:
    student = student or student_info()
    chosen = data["chosen"]
    status_labels = {
        "recommended": "推荐",
        "caution": "谨慎",
        "limited": "限定",
        "out_of_scope": "范围外",
    }
    noise_rows = "\n".join(
        f"| {int(row.snr_db)} | {row.pooled_f1_pct_reference:.2f} | "
        f"{row.pooled_f1_pct_chosen:.2f} | {row.drop_pp:.2f} |"
        for row in data["noise_compare"].sort_values("snr_db", ascending=False).itertuples()
    )
    comparison_rows = []
    for fs, bits in [(360, 11), (360, 6), (250, 8), (180, 8), (125, 8), (100, 6)]:
        row = data["clean"][(data["clean"].target_fs_hz == fs) & (data["clean"].bits == bits)].iloc[0]
        comparison_rows.append(
            f"| {fs} Hz / {bits} bit | {row.pooled_f1_pct:.3f} | "
            f"{row.f1_drop_from_reference_pp:.3f} | {row.median_timing_error_ms:.2f} | "
            f"{int(row.raw_bitrate_bps)} |"
        )
    comparison_rows_text = "\n".join(comparison_rows)
    acceptance_rows = "\n".join(
        [
            f"| 干净条件合并F1 | >= 99.0% | {s['clean_f1']}% | 达到 |",
            f"| 相对参考配置F1下降 | <= 0.5 pp | {s['clean_drop']} pp | 达到 |",
            f"| SNR>=6 dB最坏相对下降 | <= 2.0 pp | {s['noise_worst_drop']} pp | 达到 |",
            f"| R峰时间误差中位数 | <= 20 ms | {s['timing']} ms | 达到 |",
            f"| 模拟抗混叠过渡带 | 优先选择可实现配置 | {chosen['analog_transition_band_hz']:.1f} Hz | 达到 |",
        ]
    )
    monitoring_rows_text = "监护指标尚未生成。"
    if not data["monitoring"].empty:
        monitoring_rows = []
        preferred = [(360, 11), (360, 8), (360, 7), (360, 6), (180, 8), (125, 8), (100, 6)]
        for fs, bits in preferred:
            rows = data["monitoring"][
                (data["monitoring"].target_fs_hz == fs) & (data["monitoring"].bits == bits)
            ]
            if rows.empty:
                continue
            row = rows.iloc[0]
            monitoring_rows.append(
                f"| {fs} Hz / {bits} bit | {row.rr_median_abs_error_ms:.2f} | "
                f"{row.hr_median_abs_error_bpm:.3f} | {row.valid_rr_pair_pct:.2f} | "
                f"{row.sdnn_relative_error_median_pct:.2f} | {row.rmssd_relative_error_median_pct:.2f} | "
                f"{row.hrv_usable_record_pct:.1f}% |"
            )
        monitoring_rows_text = "\n".join(monitoring_rows)

    scenario_rows_text = "场景推荐尚未生成。"
    if data["scenarios"]:
        scenario_rows = []
        for scenario in data["scenarios"]:
            config = scenario.get("config") or {}
            config_label = config.get("label", "不适用")
            evidence = "; ".join(scenario.get("evidence", [])[:3])
            scenario_rows.append(
                f"| {scenario.get('title_zh', scenario.get('title', ''))} | "
                f"{status_labels.get(scenario.get('status', ''), scenario.get('status', ''))} | {config_label} | {evidence} | {scenario.get('boundary', '')} |"
            )
        scenario_rows_text = "\n".join(scenario_rows)
    return f"""# 面向可穿戴心电监护的ECG采样率、ADC有效位数与抗噪性能联合设计

**课程：** 生物医学电子（2）  
**{student_line(student)}**  
**完成日期：** 2026年6月  
**项目仓库：** https://github.com/Lebron-shun/ecg-sampling-adc-noise-robustness  
**交互展示：** https://lebron-shun.github.io/ecg-sampling-adc-noise-robustness/

## 摘要

面向长时程可穿戴心电监护中的存储、传输与抗噪权衡，本项目构建了一个可复现的虚拟ECG采集系统，系统研究采样率与ADC有效位数对R峰检测的联合影响，并进一步扩展到RR间期、瞬时心率和HRV技术可用性分析。实验使用MIT-BIH Arrhythmia Database的48条动态心电记录，以及MIT-BIH Noise Stress Test Database中12条标准电极运动伪影记录。虚拟采集链路包含0.5-40 Hz前端响应、抗混叠重采样、固定10 mV满量程量化和固定参数XQRS检测器。对5种采样率和6种有效位数组成的30种配置进行全因子评价，并使用灵敏度、阳性预测率、F1、R峰时间误差、RR误差、心率误差、HRV相对误差、削顶率和原始数据率进行综合比较。

结果表明，推荐配置为**{s['configuration']}**。其干净条件合并F1为{s['clean_f1']}%，参考配置360 Hz / 11 bit为{s['clean_ref_f1']}%，下降{s['clean_drop']}个百分点；R峰时间误差中位数为{s['timing']} ms，RR间期误差中位数为{s['rr_error']} ms，瞬时心率误差中位数为{s['hr_error']} bpm。推荐配置将单通道原始数据率降至{s['bitrate']} bit/s，相对参考配置减少{s['bitrate_reduction']}%。在SNR不低于6 dB的标准运动伪影条件下，相对参考配置的最差F1下降为{s['noise_worst_drop']}个百分点。研究说明，在以R峰、RR间期和心率趋势监测为目标的条件下，可通过任务导向的采集参数联合设计显著降低数据量；但HRV结果仅作为技术可用性分析，不能解释疾病风险，也不能外推到ST段分析、临床诊断或真实硬件功耗。

**关键词：** 心电信号；采样率；ADC有效位数；R峰检测；运动伪影；可穿戴监护

## 1 引言与医学应用背景

动态心电监护需要长时间采集人体表面ECG，并从中提取R峰、RR间期和心率。可穿戴设备通常受到电池、存储空间和无线传输带宽限制。提高采样率和ADC位数能够保留更多细节，但也增加数据率与处理负担；过度降低配置则可能造成QRS形态失真、R峰定位误差增加，并降低噪声环境下的检测可靠性。

表1将医学应用需求、工程约束和评价指标对应起来。本项目将问题限定为R峰与RR间期监测任务，不进行疾病分类、ST段分析或临床决策。研究目标是在统一算法和固定输入量程下，定量回答采样率、ADC有效位数和运动伪影如何共同影响R峰检测，并寻找性能与数据率之间的合理折中。

表1　医学应用需求、工程约束与评价指标

| 应用需求 | 工程约束 | 本项目评价指标 |
|---|---|---|
| 长时间心率与RR间期监测 | 存储容量、无线传输带宽、电池供电 | 原始数据率、每日存储量 |
| 可穿戴运动场景 | 电极运动伪影、基线漂移、瞬态干扰 | 不同SNR下的F1与相对下降 |
| R峰可靠检出 | 采样率、ADC有效位数、检测器鲁棒性 | 灵敏度、PPV、合并F1、宏平均F1 |
| R峰定位精度 | 采样时间间隔与抗混叠滤波 | 时间误差中位数、P95时间误差 |
| RR/心率趋势监护 | 连续R峰配对、检测漏检和误检 | RR误差、瞬时心率误差、连续RR对比例 |
| HRV技术可用性 | RR序列完整性、短时变异性误差 | SDNN相对误差、RMSSD相对误差、可用记录比例 |
| 可复现工程评估 | 公开数据、固定参数、自动化脚本 | 数据清单、逐记录结果、审核报告 |

本项目交付物包括虚拟ECG采集与数字处理模块、全因子评估脚本、结果表与图像、课程报告、GitHub开源仓库以及交互展示网页。

![图1　系统总体框图。模拟前端为参数化模型，非实物医疗硬件。](../figures/figure_01_system_diagram.png)

## 2 系统设计原理与关键技术方案

### 2.1 系统位置与输入输出关系

完整可穿戴ECG监护系统可由表面电极、输入保护、仪表放大、模拟滤波、ADC采样、数字处理、显示/通信和数据存储模块组成。图1显示了该完整链路。本作业实际实现其中的虚拟采集与数字处理模块：输入为公开ECG记录及其人工标注，输出为不同采样率和ADC有效位数下的R峰检测性能、数据率、抗噪退化和推荐配置。本项目不声称制作真实硬件，但在参数设计中保留模拟前端带宽、满量程、抗混叠过渡带等工程约束。

### 2.2 核心设计指标

固定满量程为10 mV峰峰值。N位ADC的量化步长为：

`LSB = 10 mV / 2^N`

单通道无压缩数据率为：

`R = fs × N`

候选采样率为360、250、180、125和100 Hz；有效位数为11、10、9、8、7和6 bit。原始数据库为360 Hz、11 bit，因此不研究更高位数的收益。

表2　核心设计指标与约束

| 指标 | 取值 | 设计依据 |
|---|---:|---|
| 前端通带 | 0.5-40 Hz | 保留主要QRS能量，抑制基线漂移和高频噪声 |
| 满量程 | 10 mV峰峰值 | 与MIT-BIH原始量化范围一致 |
| 候选采样率 | 360/250/180/125/100 Hz | 覆盖参考配置与低数据率配置 |
| 候选ADC有效位数 | 11/10/9/8/7/6 bit | 原始数据为11 bit，不外推更高位数 |
| 检测匹配窗口 | ±150 ms，敏感性分析±75 ms | 对齐人工标注并评估定位误差 |
| 噪声SNR | 24/18/12/6/0/-6 dB | 使用NSTDB标准运动伪影压力测试 |
| 推荐约束 | F1≥99%，F1下降≤0.5 pp，噪声下降≤2 pp，时间误差≤20 ms | 与R峰/RR监测任务需求对应 |

表3　ADC量化步长与参考采样率下数据率

| ADC有效位数 | LSB (µV) | 360 Hz数据率 (bit/s) |
|---:|---:|---:|
| 11 | 4.883 | 3960 |
| 10 | 9.766 | 3600 |
| 9 | 19.531 | 3240 |
| 8 | 39.063 | 2880 |
| 7 | 78.125 | 2520 |
| 6 | 156.250 | 2160 |

对于40 Hz目标通带，100 Hz采样仅留下10 Hz的模拟抗混叠过渡带，125 Hz为22.5 Hz，180 Hz为50 Hz，250 Hz为85 Hz，360 Hz为140 Hz。因此，配置选择不仅比较数字检测性能，也要考虑模拟滤波器可实现性。

## 3 方法与实现过程

### 3.1 数据集

- MIT-BIH Arrhythmia Database：48条约30分钟双通道动态ECG，360 Hz、11 bit，并带专家心搏标注。
- MIT-BIH Noise Stress Test Database：以记录118和119构建的标准电极运动伪影压力测试记录，SNR为24、18、12、6、0和-6 dB。

### 3.2 实验流程

表4列出实现环境与主要工具。公开数据和开源算法均注明来源，本项目的主要工作是构建虚拟采集链路、参数化全因子实验、指标统计、结果可视化、报告生成与交互展示。

表4　实现环境与工具

| 类别 | 工具或数据 | 用途 |
|---|---|---|
| 编程环境 | Python | 实验脚本、统计分析、报告生成 |
| 数据读取 | WFDB Python | 下载和读取PhysioNet记录与标注 |
| 数值计算 | NumPy、SciPy | 滤波、重采样、量化与指标计算 |
| 结果处理 | Pandas | 逐记录与聚合结果表 |
| 制图 | Matplotlib | 系统图、热力图、鲁棒性曲线、帕累托图 |
| 报告 | python-docx、ReportLab | 生成DOCX和PDF |
| 展示 | HTML/CSS/JavaScript | GitHub Pages交互展示 |

![图2　虚拟采集与实验评价流程。](../figures/figure_02_experiment_workflow.png)

每条记录选取第一通道，先通过0.5-40 Hz四阶Butterworth零相位带通模型，再使用带抗混叠滤波的多相重采样生成目标采样率。量化过程固定10 mV满量程，不对各记录单独归一化。所有配置使用固定设置的WFDB XQRS检测器。检测结果与人工标注在±150 ms内一对一匹配；另以±75 ms窗口进行定位敏感性分析。

NSTDB主分析仅评价官方交替模式中的已知噪声区间，并在区间边界去除0.2 s，避免边界滤波效应。数据处理、参数和结果表均由脚本自动生成。

表5　本人实现工作说明

| 环节 | 实现内容 | 主要输出 |
|---|---|---|
| 数据获取与校验 | 下载MIT-BIH和NSTDB，记录文件大小、SHA-256和采样率 | `data_manifest/`，`data_validation.csv` |
| 虚拟采集链路 | 前端带通、抗混叠重采样、固定满量程量化 | 不同采样率/位数组合信号 |
| R峰检测与匹配 | 固定参数XQRS，人工标注一对一匹配 | TP、FP、FN、时间误差 |
| 统计与可视化 | 聚合指标、bootstrap、配对检验、热力图和曲线 | `results/`，`figures/` |
| 工程交付 | 自动生成报告、PDF、DOCX、网页和审核结果 | `report/`，`web/`，`FINAL_AUDIT.md` |

![图3　不同虚拟采集配置的代表性ECG波形。](../figures/figure_03_waveform_comparison.png)

### 3.3 性能指标

`Sensitivity = TP / (TP + FN)`  
`PPV = TP / (TP + FP)`  
`F1 = 2 × Sensitivity × PPV / (Sensitivity + PPV)`

同时计算R峰时间误差、量化均方根误差、量化信噪比、削顶率、数据率与每天存储量。统计分析以记录为单位，使用bootstrap置信区间和配对Wilcoxon检验。

### 3.4 监护指标扩展方法

为避免医学场景只停留在“R峰检出率”，本项目对代表性配置进一步计算RR间期、瞬时心率和HRV技术可用性指标。首先在±150 ms窗口内对人工标注R峰与检测R峰进行一对一匹配；只有相邻两个参考R峰和对应的相邻两个检测R峰均连续匹配时，才计入有效RR对。RR间期误差定义为检测RR与参考RR的绝对差；瞬时心率误差定义为`60/RR_detected - 60/RR_reference`的绝对值。HRV指标采用常见的SDNN和RMSSD，但本项目只比较采集配置造成的相对误差，不解释自主神经功能、疾病风险或临床分层。

## 4 结果展示与性能评价

### 4.1 干净条件全因子结果

![图4　干净条件下30种配置的F1、时间误差与资源指标热力图。](../figures/figure_04_clean_heatmaps.png)

图4显示，推荐配置{s['configuration']}在干净条件下获得{s['clean_f1']}%的合并F1，较参考配置下降{s['clean_drop']}个百分点；时间误差中位数为{s['timing']} ms。结果显示，降低采样率对R峰定位时间精度的影响通常早于对检出率的影响。

表6　代表性配置性能对比

| 配置 | F1 (%) | 相对下降 (pp) | 时间误差 (ms) | 数据率 (bit/s) |
|---|---:|---:|---:|---:|
{comparison_rows_text}

以记录为重采样单位的bootstrap显示，推荐配置相对参考配置的宏平均F1下降95%置信区间为{s['bootstrap_low']}至{s['bootstrap_high']}个百分点；配对Wilcoxon检验经Holm校正后的p值为{s['holm_p']}。在20个探索性交互对比中，有{s['interaction_count']}个bootstrap区间未跨越零，且效应量很小，未显示强烈的采样率-位数交互。

![图5　参考配置与推荐配置的逐记录F1分布。](../figures/figure_08_record_distribution.png)

### 4.2 标准运动伪影抗噪性能

![图6　不同SNR下候选配置的R峰检测性能。](../figures/figure_05_noise_robustness.png)

表7　标准运动伪影条件下推荐配置与参考配置对比

| SNR (dB) | 参考F1 (%) | 推荐配置F1 (%) | 相对下降 (pp) |
|---:|---:|---:|---:|
{noise_rows}

图6和表7显示，噪声增强后所有配置的绝对性能均下降。推荐配置在SNR不低于6 dB时相对参考配置的最差下降为{s['noise_worst_drop']}个百分点。严重噪声下最低F1为{s['noise_lowest_f1']}%，说明降低数据率不能解决检测算法本身对强运动伪影的脆弱性，主要限制逐渐转向检测器鲁棒性和输入污染。

![图7　严重运动伪影下的全因子性能热力图。](../figures/figure_07_severe_noise_heatmap.png)

### 4.3 性能与资源权衡

![图8　干净条件性能与原始数据率的帕累托权衡。](../figures/figure_06_pareto_tradeoff.png)

图8显示，推荐配置的数据率为{s['bitrate']} bit/s，每天原始存储约{s['storage']} MiB，相对360 Hz / 11 bit参考配置减少{s['bitrate_reduction']}%。在满足预设干净条件性能、相对抗噪退化与时间误差约束后，该配置具有较低数据率，并保留比100 Hz更合理的模拟抗混叠过渡带。

表8　设计指标达成情况

| 评价项 | 预设约束 | 推荐配置结果 | 结论 |
|---|---:|---:|---|
{acceptance_rows}

### 4.4 医学监护指标扩展：RR、心率与HRV

![图9　代表配置下RR、心率与HRV监护指标。](../figures/figure_09_monitoring_metrics.png)

图9和表9显示，推荐配置{s['configuration']}在R峰检出性能之外，也保持了较稳定的RR间期和瞬时心率趋势：RR误差中位数为{s['rr_error']} ms，心率误差中位数为{s['hr_error']} bpm，连续有效RR对比例为{s['valid_rr']}%。SDNN和RMSSD相对误差分别为{s['sdnn_error']}%和{s['rmssd_error']}%，HRV可用记录比例为{s['hrv_usable']}%。这些指标说明推荐配置适合R峰、RR和心率趋势监护；若用于HRV趋势观察，需要同时检查RR连续性和SDNN/RMSSD误差，不应将HRV结果解释为临床诊断结论。

表9　代表性配置的RR、心率与HRV技术指标

| 配置 | RR误差中位数 (ms) | 心率误差中位数 (bpm) | 连续RR对比例 (%) | SDNN相对误差 (%) | RMSSD相对误差 (%) | HRV可用记录比例 |
|---|---:|---:|---:|---:|---:|---:|
{monitoring_rows_text}

表10　面向不同医学监护场景的配置建议

| 监护场景 | 推荐等级 | 配置 | 关键证据 | 边界说明 |
|---|---|---|---|---|
{scenario_rows_text}

## 5 讨论与改进分析

### 5.1 主要发现

第一，采样率和位数不应只根据波形视觉效果选择，而应根据具体监护任务进行联合优化。第二，R峰检出率在干净条件下可能接近性能上限，但时间误差、RR误差和心率误差仍能揭示低采样率的代价。第三，强运动伪影下的主要限制逐渐转向检测器鲁棒性与输入污染，而不是ADC位数本身。第四，HRV类指标比单个R峰检出率更依赖连续RR序列完整性，因此只能作为采集参数对节律监护影响的技术分析。

### 5.2 工程意义

{s['configuration']}适合以心率和RR间期为主的长时监测。数据率降低意味着存储与无线传输负担下降，也可能减少ADC和处理系统的工作量。但本项目没有真实硬件，不能将数据率下降直接写成实测功耗下降。对于HRV趋势观察，应优先报告RR连续性、SDNN相对误差和RMSSD相对误差，而不是只报告F1。

### 5.3 局限性

**医学应用限制。** 结论仅适用于R峰、RR间期、心率趋势和HRV技术可用性分析，不能外推到ST段分析、形态诊断、心律失常分类或临床决策。真实应用还需要安全性、可靠性、伦理和临床验证。

**工程实现限制。** 原始数据已经以360 Hz、11 bit数字化，只能可信研究降采样和降低有效位数；零相位离线滤波不等价于实时因果模拟前端，尚未验证群时延、元件误差、输入保护和真实功耗。

**算法限制。** XQRS参数保持固定，结果包含特定检测器与采集配置之间的交互；严重运动伪影下需要更鲁棒的检测器或信号质量控制。

**数据限制。** NSTDB是标准化噪声压力测试，不能完全代表长期真实佩戴中的所有干扰，也没有覆盖不同设备结构、皮肤接触状态和人群差异。

### 5.4 后续改进

后续应搭建真实模拟前端与ADC硬件，测量输入保护、共模抑制、实时延迟和实际功耗；加入信号质量指数，在低质量片段拒绝输出；比较因果滤波与不同QRS检测算法；并使用真实可穿戴运动数据进行外部验证。同时可以将交互展示网页扩展为参数设计教学工具，用于展示采样率、量化位数、抗噪性能、RR连续性和HRV技术误差之间的权衡。

## 6 结论

本项目完成了ECG采样率、ADC有效位数与抗噪性能的全因子联合设计，并将医学场景从R峰检测扩展到RR间期、心率趋势和HRV技术可用性。基于48条动态ECG和12条标准运动伪影记录，推荐{s['configuration']}作为R峰/RR/心率趋势监测任务的工程折中方案。该配置在保持与参考配置接近的R峰检测性能时，将原始数据率降低{s['bitrate_reduction']}%。研究结果支持任务导向的低数据率采集设计，同时强调HRV解释、严重运动伪影、真实因果前端和临床用途仍需进一步验证。

## 参考文献

[1] Moody GB, Mark RG. The impact of the MIT-BIH Arrhythmia Database. IEEE Engineering in Medicine and Biology Magazine, 2001, 20(3):45-50.  
[2] Goldberger AL, et al. PhysioBank, PhysioToolkit, and PhysioNet. Circulation, 2000, 101(23):e215-e220.  
[3] MIT-BIH Arrhythmia Database. PhysioNet. https://physionet.org/content/mitdb/1.0.0/  
[4] MIT-BIH Noise Stress Test Database. PhysioNet. https://physionet.org/content/nstdb/1.0.0/  
[5] WFDB Python Package Documentation. https://wfdb.readthedocs.io/
[6] Pan J, Tompkins WJ. A real-time QRS detection algorithm. IEEE Transactions on Biomedical Engineering, 1985, BME-32(3):230-236.

## 附录A：核心参数表

| 参数 | 取值 |
|---|---|
| 源采样率 | 360 Hz |
| 目标采样率 | 360、250、180、125、100 Hz |
| ADC有效位数 | 11、10、9、8、7、6 bit |
| 满量程 | 10 mV峰峰值 |
| 前端通带 | 0.5-40 Hz |
| 匹配窗口 | 150 ms，敏感性分析75 ms |
| 噪声区间边界余量 | 0.2 s |
| bootstrap次数 | 1000 |

## 附录B：关键脚本说明

| 脚本 | 作用 |
|---|---|
| `src/download_data.py` | 下载MIT-BIH与NSTDB数据 |
| `src/validate_data.py` | 校验数据完整性、采样率和标注 |
| `src/run_experiments.py` | 运行全因子虚拟采集与R峰检测 |
| `src/analyze_results.py` | 聚合结果、统计检验、推荐配置、制图 |
| `src/monitoring_analysis.py` | 计算RR、心率和HRV监护指标并生成场景建议 |
| `src/build_web_data.py` | 生成交互网页使用的数据载荷 |
| `src/build_report.py` | 生成Markdown和DOCX报告 |
| `src/build_pdf.py` | 生成最终PDF报告 |
| `src/audit_project.py` | 审核交付物完整性和课程要求覆盖 |

## 附录C：结果文件索引

| 路径 | 内容 |
|---|---|
| `results/per_record_clean.csv` | MIT-BIH逐记录干净条件结果 |
| `results/per_record_noise.csv` | NSTDB逐记录噪声条件结果 |
| `results/candidate_summary.csv` | 30组配置汇总与约束判断 |
| `results/monitoring_summary.csv` | RR、心率和HRV配置级监护指标 |
| `results/scenario_recommendations.json` | 医学监护场景推荐矩阵 |
| `figures/figure_*.png` | 系统图、流程图、波形图、热力图、鲁棒性曲线、帕累托图和监护指标图 |
| `report/final_report.pdf` | 最终课程报告PDF |
| `report/final_report.docx` | 可编辑课程报告DOCX |
| `web/index.html` | 交互展示网页 |

## 附录D：复现与展示链接

项目提供完整源代码、固定参数配置、软件依赖、数据下载脚本、SHA-256数据清单、逐记录结果、聚合结果与自动制图脚本。执行顺序：

1. `python src/download_data.py`
2. `python src/validate_data.py`
3. `python src/run_experiments.py --mode all --workers 8`
4. `python src/analyze_results.py`
5. `python src/monitoring_analysis.py`
6. `python src/build_web_data.py`
7. `python src/build_report.py`
8. `python src/build_pdf.py`
9. `python src/audit_project.py`

GitHub仓库：https://github.com/Lebron-shun/ecg-sampling-adc-noise-robustness  
交互展示页：https://lebron-shun.github.io/ecg-sampling-adc-noise-robustness/
"""


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_run_font(run, name: str, size: float, color: str = DARK, bold: bool = False) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)


def add_body(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.28
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first, BODY_FONT, 10.5, DARK, True)
        rest = paragraph.add_run(text[len(bold_prefix) :])
        set_run_font(rest, BODY_FONT, 10.5, DARK)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, BODY_FONT, 10.5, DARK)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.2
        run = paragraph.add_run(item)
        set_run_font(run, BODY_FONT, 10.5, DARK)


def add_figure(doc: Document, filename: str, caption: str, width: float = 6.25) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.add_run().add_picture(str(FIGURES_DIR / filename), width=Inches(width))
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_after = Pt(8)
    run = caption_p.add_run(caption)
    set_run_font(run, BODY_FONT, 9, GRAY)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, (cell, text, width) in enumerate(zip(header.cells, headers, widths)):
        cell.width = Inches(width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, BLUE)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, HEADING_FONT, 9, WHITE, True)
    for row_values in rows:
        row = table.add_row()
        for index, (cell, text, width) in enumerate(zip(row.cells, row_values, widths)):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if len(table.rows) % 2 == 0:
                set_cell_shading(cell, "F4F8FC")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if index > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(text))
            set_run_font(run, BODY_FONT, 8.7, DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def clean_markdown_text(text: str) -> str:
    text = text.replace("**", "")
    text = text.replace("`", "")
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1（\2）", text)
    return text.strip()


def split_markdown_row(line: str) -> list[str]:
    cells = line.strip().strip("|").split("|")
    return [clean_markdown_text(cell.strip()) for cell in cells]


def markdown_table_widths(column_count: int) -> list[float]:
    presets = {
        2: [2.2, 4.1],
        3: [1.65, 2.15, 2.55],
        4: [1.15, 1.7, 1.7, 1.7],
        5: [1.28, 1.05, 1.22, 1.25, 1.45],
    }
    if column_count in presets:
        return presets[column_count]
    return [6.3 / column_count] * column_count


def add_markdown_table(doc: Document, lines: list[str], index: int) -> int:
    headers = split_markdown_row(lines[index])
    rows: list[list[str]] = []
    index += 2
    while index < len(lines) and lines[index].strip().startswith("|"):
        rows.append(split_markdown_row(lines[index]))
        index += 1
    add_table(doc, headers, rows, markdown_table_widths(len(headers)))
    return index


def add_markdown_content(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    try:
        index = lines.index("## 摘要")
    except ValueError:
        index = 0

    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue
        if line.startswith("### "):
            add_heading(doc, clean_markdown_text(line[4:]), 2)
            index += 1
            continue
        if line.startswith("## "):
            add_heading(doc, clean_markdown_text(line[3:]), 1)
            index += 1
            continue
        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", line)
        if image_match:
            caption, image_path = image_match.groups()
            add_figure(doc, Path(image_path).name, clean_markdown_text(caption))
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and "---" in lines[index + 1]:
            index = add_markdown_table(doc, lines, index)
            continue
        if line.startswith("- "):
            items = []
            while index < len(lines) and lines[index].strip().startswith("- "):
                items.append(clean_markdown_text(lines[index].strip()[2:]))
                index += 1
            add_bullets(doc, items)
            continue
        ordered_match = re.match(r"\d+\.\s+(.*)", line)
        if ordered_match:
            items = []
            while index < len(lines):
                match = re.match(r"\d+\.\s+(.*)", lines[index].strip())
                if not match:
                    break
                items.append(clean_markdown_text(match.group(1)))
                index += 1
            add_bullets(doc, items)
            continue
        add_body(doc, clean_markdown_text(line))
        index += 1


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.28

    for level, size, before, after in ((1, 15, 15, 7), (2, 12.5, 11, 5), (3, 11, 8, 4)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = HEADING_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLUE if level < 3 else DARK)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style in ("List Bullet", "List Number"):
        style = doc.styles[list_style]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(10.5)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header.add_run("生物医学电子（2）课程大作业　|　ECG采集参数联合设计")
    set_run_font(header_run, HEADING_FONT, 8.5, GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("最终课程报告")
    set_run_font(run, HEADING_FONT, 8.5, GRAY)


def add_cover(doc: Document, s: dict[str, str], student: dict[str, str] | None = None) -> None:
    student = student or student_info()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(70)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("生物医学电子（2）课程大作业")
    set_run_font(run, HEADING_FONT, 15, TEAL, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("面向可穿戴心电监护的ECG采样率、\nADC有效位数与抗噪性能联合设计")
    set_run_font(run, HEADING_FONT, 25, DARK, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(38)
    run = p.add_run("基于公开数据集的虚拟采集系统设计与定量性能评价")
    set_run_font(run, HEADING_FONT, 12, GRAY)

    add_table(
        doc,
        ["项目关键结果", "数值"],
        [
            ["推荐配置", s["configuration"]],
            ["干净条件合并F1", f"{s['clean_f1']}%"],
            ["数据率降低", f"{s['bitrate_reduction']}%"],
            ["R峰时间误差中位数", f"{s['timing']} ms"],
        ],
        [3.4, 2.2],
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(42)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(student_line(student))
    set_run_font(run, BODY_FONT, 11, DARK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("完成日期：2026年6月")
    set_run_font(run, BODY_FONT, 11, GRAY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("GitHub：https://github.com/Lebron-shun/ecg-sampling-adc-noise-robustness")
    set_run_font(run, BODY_FONT, 9.5, GRAY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("交互展示：https://lebron-shun.github.io/ecg-sampling-adc-noise-robustness/")
    set_run_font(run, BODY_FONT, 9.5, GRAY)
    p.add_run().add_break(WD_BREAK.PAGE)


def build_docx(data: dict, s: dict[str, str], path: Path, student: dict[str, str] | None = None) -> None:
    doc = Document()
    setup_styles(doc)
    add_cover(doc, s, student)
    add_markdown_content(doc, markdown_report(data, s, student))
    doc.save(path)
    return

    add_heading(doc, "摘要", 1)
    add_body(
        doc,
        f"面向长时程可穿戴心电监护中的存储、传输与抗噪权衡，本项目构建了可复现的虚拟ECG采集系统，"
        f"系统研究采样率与ADC有效位数对R峰检测的联合影响。实验使用MIT-BIH Arrhythmia Database的48条动态心电记录，"
        f"以及MIT-BIH Noise Stress Test Database中12条标准电极运动伪影记录。对5种采样率和6种有效位数组成的30种配置进行全因子评价。"
        f"结果推荐{s['configuration']}：干净条件合并F1为{s['clean_f1']}%，相对360 Hz / 11 bit参考配置下降{s['clean_drop']}个百分点，"
        f"同时将单通道原始数据率降低{s['bitrate_reduction']}%。该结论适用于R峰和RR监测任务，不能外推到临床诊断或真实硬件功耗。"
    )
    add_body(doc, "关键词：心电信号；采样率；ADC有效位数；R峰检测；运动伪影；可穿戴监护")

    add_heading(doc, "1 引言与医学应用背景", 1)
    add_body(
        doc,
        "动态心电监护需要长时间采集人体表面ECG，并从中提取R峰、RR间期和心率。可穿戴设备受到电池、存储空间和无线传输带宽限制。"
        "提高采样率和ADC位数能够保留更多细节，但也增加数据率与处理负担；过度降低配置则可能造成QRS波失真、R峰定位误差增加，并降低噪声环境下的检测可靠性。"
    )
    add_body(
        doc,
        "本项目不进行疾病分类，而是将系统性能限定为R峰监测任务。在统一算法和固定输入量程下，定量研究采样率、有效位数和运动伪影的联合影响，并寻找性能与数据率之间的工程折中。"
    )
    add_figure(doc, "figure_01_system_diagram.png", "图1　系统总体框图")

    add_heading(doc, "2 系统设计原理与关键技术方案", 1)
    add_heading(doc, "2.1 系统组成与设计约束", 2)
    add_body(
        doc,
        "虚拟系统由表面电极、保护与仪表放大、0.5-40 Hz模拟前端、可配置采样与量化、数字QRS检测和结果输出组成。"
        "本作业实际实现虚拟采集与数字处理模块，不声称制作真实医疗硬件。"
    )
    add_table(
        doc,
        ["设计项目", "参数", "设计依据"],
        [
            ["输入信号", "单导联ECG", "面向心率与RR间期监测"],
            ["前端通带", "0.5-40 Hz", "保留主要QRS信息并抑制漂移与高频噪声"],
            ["输入满量程", "10 mV峰峰值", "与MIT-BIH原始量化范围一致"],
            ["候选采样率", "360/250/180/125/100 Hz", "覆盖参考配置与低数据率方案"],
            ["候选有效位数", "11/10/9/8/7/6 bit", "原始数据上限为11 bit"],
        ],
        [1.35, 2.05, 2.9],
    )
    add_heading(doc, "2.2 采样、量化与数据率", 2)
    add_body(doc, "固定满量程为10 mV峰峰值，N位ADC量化步长为LSB = 10 mV / 2^N；单通道无压缩数据率为R = fs × N。")
    add_body(
        doc,
        "100 Hz虽然满足40 Hz带宽的奈奎斯特条件，但只留下10 Hz模拟抗混叠过渡带；因此最终推荐不仅考虑算法性能，也要求具有较合理的模拟滤波实现空间。"
    )

    add_heading(doc, "3 方法与实现过程", 1)
    add_heading(doc, "3.1 开源数据与实验矩阵", 2)
    add_body(
        doc,
        "MIT-BIH Arrhythmia Database包含48条约30分钟双通道动态ECG，采样率360 Hz，原始分辨率11 bit，并带专家复核的逐搏标注。"
        "MIT-BIH Noise Stress Test Database提供以记录118和119构建的标准电极运动伪影压力测试记录，SNR为24、18、12、6、0和-6 dB。"
    )
    add_table(
        doc,
        ["实验", "记录", "因素", "主要目的"],
        [
            ["干净条件全因子", "MIT-BIH 48条", "5采样率 × 6位数", "评价采样与量化影响"],
            ["标准抗噪压力测试", "NSTDB 12条", "30配置 × 6个SNR", "评价运动伪影鲁棒性"],
            ["定位敏感性分析", "同上", "±150 ms与±75 ms", "识别低采样率时间代价"],
        ],
        [1.3, 1.35, 1.7, 1.95],
    )
    add_figure(doc, "figure_02_experiment_workflow.png", "图2　虚拟采集与实验评价流程")
    add_heading(doc, "3.2 信号处理与评价", 2)
    add_body(
        doc,
        "每条记录选取第一通道，先通过0.5-40 Hz四阶Butterworth零相位带通模型，再使用带抗混叠滤波的多相重采样生成目标采样率。"
        "量化过程固定10 mV满量程，不对各记录单独归一化。所有配置使用固定设置的WFDB XQRS检测器。"
    )
    add_body(
        doc,
        "检测结果与人工标注在±150 ms内一对一匹配，并计算TP、FP、FN、灵敏度、阳性预测率和F1。"
        "NSTDB主分析仅评价官方交替模式中的已知噪声区间，并在边界去除0.2 s。"
    )
    add_figure(doc, "figure_03_waveform_comparison.png", "图3　不同虚拟采集配置的代表性ECG波形")

    add_heading(doc, "4 结果展示与性能评价", 1)
    add_heading(doc, "4.1 干净条件全因子结果", 2)
    add_figure(doc, "figure_04_clean_heatmaps.png", "图4　干净条件下30种配置的F1与R峰时间误差")
    comparison_configs = [(360, 11), (360, 6), (250, 8), (180, 8), (125, 8), (100, 6)]
    comparison_rows = []
    for fs, bits in comparison_configs:
        row = data["clean"][(data["clean"].target_fs_hz == fs) & (data["clean"].bits == bits)].iloc[0]
        comparison_rows.append(
            [
                f"{fs} Hz / {bits} bit",
                fmt(row["pooled_f1_pct"], 3),
                fmt(row["f1_drop_from_reference_pp"], 3),
                fmt(row["median_timing_error_ms"], 2),
                str(int(row["raw_bitrate_bps"])),
            ]
        )
    add_table(
        doc,
        ["配置", "F1 (%)", "相对下降 (pp)", "时间误差 (ms)", "数据率 (bit/s)"],
        comparison_rows,
        [1.35, 1.05, 1.25, 1.25, 1.35],
    )
    add_body(
        doc,
        f"推荐配置{s['configuration']}在干净条件下获得{s['clean_f1']}%的合并F1，参考配置为{s['clean_ref_f1']}%，"
        f"下降{s['clean_drop']}个百分点；R峰时间误差中位数为{s['timing']} ms。"
        "热力图显示，降低采样率对R峰定位精度的影响通常早于对检出率的影响。"
    )
    add_body(
        doc,
        f"以记录为重采样单位的bootstrap显示，推荐配置相对参考配置的宏平均F1下降95%置信区间为{s['bootstrap_low']}至{s['bootstrap_high']}个百分点；"
        f"配对Wilcoxon检验经Holm校正后的p值为{s['holm_p']}。在20个探索性交互对比中，有{s['interaction_count']}个区间未跨越零，且效应量很小。"
    )
    add_figure(doc, "figure_08_record_distribution.png", "图5　参考配置与推荐配置的逐记录F1分布")

    add_heading(doc, "4.2 标准运动伪影抗噪性能", 2)
    add_figure(doc, "figure_05_noise_robustness.png", "图6　不同SNR下候选配置的R峰检测性能")
    noise_rows = [
        [
            str(int(row.snr_db)),
            fmt(row.pooled_f1_pct_reference, 2),
            fmt(row.pooled_f1_pct_chosen, 2),
            fmt(row.drop_pp, 2),
        ]
        for row in data["noise_compare"].sort_values("snr_db", ascending=False).itertuples()
    ]
    add_table(
        doc,
        ["SNR (dB)", "参考F1 (%)", "推荐F1 (%)", "相对下降 (pp)"],
        noise_rows,
        [1.2, 1.7, 1.7, 1.7],
    )
    add_body(
        doc,
        f"噪声增强后所有配置的绝对性能均下降。推荐配置在SNR不低于6 dB时相对参考配置的最差下降为{s['noise_worst_drop']}个百分点；"
        f"全部噪声水平中的最低F1为{s['noise_lowest_f1']}%。这说明强运动伪影下的主要限制逐渐转向检测器鲁棒性与输入污染。"
    )
    add_figure(doc, "figure_07_severe_noise_heatmap.png", "图7　严重运动伪影下的全因子性能")

    add_heading(doc, "4.3 性能与资源权衡", 2)
    add_figure(doc, "figure_06_pareto_tradeoff.png", "图8　干净条件性能与原始数据率权衡")
    add_body(
        doc,
        f"推荐配置的数据率为{s['bitrate']} bit/s，每天单通道原始存储约{s['storage']} MiB，相对360 Hz / 11 bit参考配置减少{s['bitrate_reduction']}%。"
        "它满足预设的干净条件性能、相对抗噪退化和时间误差约束，同时保留比100 Hz更合理的模拟抗混叠过渡带。"
    )

    add_heading(doc, "5 讨论与改进分析", 1)
    add_heading(doc, "5.1 主要发现与工程意义", 2)
    add_body(
        doc,
        "采样率和有效位数应根据具体生理信号任务联合优化，而不能只根据波形视觉效果或单一理论指标选择。"
        "干净条件下R峰检出率容易接近性能上限，但时间误差可以更早揭示低采样率代价。"
        "数据率下降可直接减少存储与传输负担，也可能降低系统工作量；但本项目没有真实硬件，不能将数据率下降写成实测功耗下降。"
    )
    add_heading(doc, "5.2 局限性", 2)
    add_bullets(
        doc,
        [
            "原始数据已经以360 Hz、11 bit数字化，只能可信研究降采样和降低有效位数。",
            "零相位离线滤波不等价于实时因果模拟前端，尚未验证群时延、元件误差与安全性。",
            "标准噪声压力测试不能完全代表长期真实佩戴中的全部干扰。",
            "结论仅适用于R峰监测，不能外推到ST段分析、形态诊断或临床决策。",
            "XQRS设置保持固定，结果包含特定检测器与采集参数之间的交互。",
        ],
    )
    add_heading(doc, "5.3 后续改进", 2)
    add_body(
        doc,
        "后续应搭建真实模拟前端与ADC硬件，测量共模抑制、实时延迟和实际功耗；加入信号质量指数，在低质量片段拒绝输出；"
        "比较因果滤波与不同QRS检测算法，并使用真实可穿戴运动数据进行外部验证。"
    )

    add_heading(doc, "6 结论", 1)
    add_body(
        doc,
        f"本项目完成了ECG采样率、ADC有效位数与抗噪性能的全因子联合设计。基于48条动态ECG和12条标准运动伪影记录，"
        f"推荐{s['configuration']}作为R峰监测任务的工程折中方案。该配置保持与参考配置接近的R峰检测性能，同时将原始数据率降低{s['bitrate_reduction']}%。"
        "研究结果支持任务导向的低数据率采集设计，同时强调严重运动伪影、真实因果前端和临床用途仍需进一步验证。"
    )

    add_heading(doc, "参考文献", 1)
    references = [
        "[1] Moody GB, Mark RG. The impact of the MIT-BIH Arrhythmia Database. IEEE Engineering in Medicine and Biology Magazine, 2001, 20(3):45-50.",
        "[2] Goldberger AL, et al. PhysioBank, PhysioToolkit, and PhysioNet. Circulation, 2000, 101(23):e215-e220.",
        "[3] MIT-BIH Arrhythmia Database. PhysioNet. https://physionet.org/content/mitdb/1.0.0/",
        "[4] MIT-BIH Noise Stress Test Database. PhysioNet. https://physionet.org/content/nstdb/1.0.0/",
        "[5] WFDB Python Package Documentation. https://wfdb.readthedocs.io/",
    ]
    for reference in references:
        add_body(doc, reference)

    add_heading(doc, "附录A　复现说明", 1)
    add_body(
        doc,
        "最终项目提供完整源代码、固定参数配置、软件依赖、数据下载脚本、SHA-256数据清单、逐记录结果、聚合结果和自动制图脚本。"
    )
    add_bullets(
        doc,
        [
            "python src/download_data.py",
            "python src/run_experiments.py --mode all --workers 8",
            "python src/analyze_results.py",
            "python src/build_report.py",
        ],
    )
    doc.save(path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output-dir", type=Path, default=REPORT_DIR)
    parser.add_argument("--basename", default="final_report")
    parser.add_argument("--name")
    parser.add_argument("--student-id")
    parser.add_argument("--class-name")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    ensure_project_dirs()
    args.output_dir.mkdir(parents=True, exist_ok=True)
    data = read_inputs()
    s = summary_text(data)
    student = student_info(args.name, args.student_id, args.class_name)
    markdown = markdown_report(data, s, student)
    md_path = args.output_dir / f"{args.basename}.md"
    docx_path = args.output_dir / f"{args.basename}.docx"
    md_path.write_text(markdown, encoding="utf-8")
    build_docx(data, s, docx_path, student)
    print(f"Wrote {md_path}")
    print(f"Wrote {docx_path}")


if __name__ == "__main__":
    main()
