"""Completion audit for the final ECG joint-design submission package."""

from __future__ import annotations

import hashlib
import json
import zipfile
from pathlib import Path

import pandas as pd
from docx import Document
from pypdf import PdfReader

from project_core import FIGURES_DIR, PROJECT_ROOT, REPORT_DIR, RESULTS_DIR


AUDIT_JSON = PROJECT_ROOT / "FINAL_AUDIT.json"
AUDIT_MD = PROJECT_ROOT / "FINAL_AUDIT.md"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        while chunk := handle.read(1024 * 1024):
            digest.update(chunk)
    return digest.hexdigest()


def check(name: str, condition: bool, evidence: str) -> dict:
    return {"requirement": name, "passed": bool(condition), "evidence": evidence}


def docx_text(document: Document) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def main() -> None:
    checks = []
    validation = json.loads((RESULTS_DIR / "data_validation_summary.json").read_text(encoding="utf-8"))
    clean = pd.read_csv(RESULTS_DIR / "per_record_clean.csv")
    noise = pd.read_csv(RESULTS_DIR / "per_record_noise.csv")
    clean_agg = pd.read_csv(RESULTS_DIR / "aggregate_clean.csv")
    noise_agg = pd.read_csv(RESULTS_DIR / "aggregate_noise_by_snr.csv")
    candidates = pd.read_csv(RESULTS_DIR / "candidate_summary.csv")
    recommended = json.loads((RESULTS_DIR / "recommended_config.json").read_text(encoding="utf-8"))
    monitoring_path = RESULTS_DIR / "monitoring_summary.csv"
    monitoring_metrics_path = RESULTS_DIR / "monitoring_metrics.csv"
    scenarios_path = RESULTS_DIR / "scenario_recommendations.json"
    monitoring = pd.read_csv(monitoring_path) if monitoring_path.exists() else pd.DataFrame()
    monitoring_metrics = pd.read_csv(monitoring_metrics_path) if monitoring_metrics_path.exists() else pd.DataFrame()
    scenarios = json.loads(scenarios_path.read_text(encoding="utf-8")) if scenarios_path.exists() else []

    checks.append(
        check(
            "Open-source data completeness",
            validation == {
                "total_records": 60,
                "mitdb_records": 48,
                "nstdb_records": 12,
                "all_360_hz": True,
                "all_nonempty": True,
                "failures": 0,
            },
            json.dumps(validation, ensure_ascii=False),
        )
    )
    checks.append(
        check(
            "Clean full-factorial coverage",
            len(clean) == 2880
            and clean.record.nunique() == 48
            and clean[["target_fs_hz", "bits"]].drop_duplicates().shape[0] == 30
            and clean[(clean.eval_scope == "clean_after_5min") & (clean.tolerance_ms == 150)]
            .groupby(["target_fs_hz", "bits"])
            .record.nunique()
            .eq(48)
            .all(),
            f"rows={len(clean)}, records={clean.record.nunique()}, configurations=30",
        )
    )
    checks.append(
        check(
            "Noise full-factorial coverage",
            len(noise) == 1440
            and noise.record.nunique() == 12
            and noise[["target_fs_hz", "bits"]].drop_duplicates().shape[0] == 30
            and len(noise_agg) == 180,
            f"rows={len(noise)}, records={noise.record.nunique()}, aggregate SNR-config rows={len(noise_agg)}",
        )
    )
    checks.append(
        check(
            "One-to-one metric count conservation",
            bool(
                ((clean.tp + clean.fn) == clean.reference_beats).all()
                and ((clean.tp + clean.fp) == clean.detected_beats).all()
                and ((noise.tp + noise.fn) == noise.reference_beats).all()
                and ((noise.tp + noise.fp) == noise.detected_beats).all()
            ),
            "All clean and noise rows satisfy TP+FN=reference beats and TP+FP=detected beats",
        )
    )
    checks.append(
        check(
            "No experiment failures",
            json.loads((RESULTS_DIR / "failures_mitdb.json").read_text()) == {}
            and json.loads((RESULTS_DIR / "failures_nstdb.json").read_text()) == {},
            "failures_mitdb.json={} and failures_nstdb.json={}",
        )
    )
    chosen_row = candidates[
        (candidates.target_fs_hz == recommended["target_fs_hz"])
        & (candidates.bits == recommended["bits"])
    ].iloc[0]
    checks.append(
        check(
            "Recommended configuration satisfies declared constraints",
            bool(chosen_row.meets_all and chosen_row.analog_practical),
            f"{int(chosen_row.target_fs_hz)} Hz/{int(chosen_row.bits)} bit; "
            f"clean F1={chosen_row.pooled_f1_pct:.3f}%; "
            f"worst >=6 dB relative noise drop={chosen_row.worst_noise_f1_drop_pp:.3f} pp",
        )
    )
    expected_figures = [FIGURES_DIR / f"figure_{index:02d}_{suffix}.png" for index, suffix in [
        (1, "system_diagram"),
        (2, "experiment_workflow"),
        (3, "waveform_comparison"),
        (4, "clean_heatmaps"),
        (5, "noise_robustness"),
        (6, "pareto_tradeoff"),
        (7, "severe_noise_heatmap"),
        (8, "record_distribution"),
        (9, "monitoring_metrics"),
    ]]
    checks.append(
        check(
            "Required result figures",
            all(path.exists() and path.stat().st_size > 20_000 for path in expected_figures),
            f"{sum(path.exists() for path in expected_figures)}/9 figures exist",
        )
    )
    checks.append(
        check(
            "Monitoring extension outputs",
            monitoring_path.exists()
            and monitoring_metrics_path.exists()
            and scenarios_path.exists()
            and len(monitoring) == 7
            and len(monitoring_metrics) == 48 * 7
            and monitoring[["target_fs_hz", "bits"]].drop_duplicates().shape[0] == 7
            and all(keyword in monitoring.columns for keyword in [
                "rr_median_abs_error_ms",
                "hr_median_abs_error_bpm",
                "valid_rr_pair_pct",
                "sdnn_relative_error_median_pct",
                "rmssd_relative_error_median_pct",
                "hrv_usable_record_pct",
            ])
            and len(scenarios) >= 5,
            f"monitoring_summary_rows={len(monitoring)}, monitoring_metric_rows={len(monitoring_metrics)}, scenarios={len(scenarios)}",
        )
    )
    checks.append(
        check(
            "Figure 9 editable export set",
            all(
                (FIGURES_DIR / f"figure_09_monitoring_metrics.{suffix}").exists()
                and (FIGURES_DIR / f"figure_09_monitoring_metrics.{suffix}").stat().st_size > 5_000
                for suffix in ["png", "svg", "pdf"]
            )
            and "<text" in (FIGURES_DIR / "figure_09_monitoring_metrics.svg").read_text(encoding="utf-8", errors="ignore"),
            "figure_09_monitoring_metrics exported as PNG/SVG/PDF with SVG text elements",
        )
    )
    pdf_path = REPORT_DIR / "final_report.pdf"
    reader = PdfReader(pdf_path)
    page_text = [(page.extract_text() or "").strip() for page in reader.pages]
    pdf_text = "\n".join(page_text)
    checks.append(
        check(
            "Final PDF report and visual content",
            len(reader.pages) >= 9
            and all(len(text) > 20 for text in page_text)
            and all(keyword in pdf_text for keyword in ["引言", "系统设计", "方法", "结果展示", "讨论", "结论", "参考文献", "附录"]),
            f"pages={len(reader.pages)}, nonempty_pages={sum(bool(text) for text in page_text)}, sha256={sha256(pdf_path)}",
        )
    )
    docx_path = REPORT_DIR / "final_report.docx"
    doc = Document(docx_path)
    editable_text = docx_text(doc)
    with zipfile.ZipFile(docx_path) as archive:
        images = [name for name in archive.namelist() if name.startswith("word/media/")]
    checks.append(
        check(
            "Editable DOCX report",
            len(doc.paragraphs) >= 90 and len(doc.tables) >= 10 and len(images) >= 9,
            f"paragraphs={len(doc.paragraphs)}, tables={len(doc.tables)}, embedded_images={len(images)}, sha256={sha256(docx_path)}",
        )
    )
    md_path = REPORT_DIR / "final_report.md"
    md_text = md_path.read_text(encoding="utf-8")
    combined_report_text = "\n".join([md_text, pdf_text, editable_text])
    placeholder_tokens = ["____________", "姓名：________", "学号：________", "班级：________"]
    checks.append(
        check(
            "No raw personal-info placeholders in generated reports",
            not any(token in combined_report_text for token in placeholder_tokens)
            and "提交前填写" in combined_report_text,
            "Generated MD/PDF/DOCX use 提交前填写 instead of underline placeholders",
        )
    )
    expected_table_labels = [f"表{index}" for index in range(1, 11)]
    expected_figure_labels = [f"图{index}" for index in range(1, 10)]
    checks.append(
        check(
            "Numbered tables and figures in report",
            all(label in combined_report_text for label in expected_table_labels + expected_figure_labels),
            f"tables={sum(label in combined_report_text for label in expected_table_labels)}/10, "
            f"figures={sum(label in combined_report_text for label in expected_figure_labels)}/9",
        )
    )
    checks.append(
        check(
            "RR/HRV monitoring methods and boundaries in report",
            all(keyword in combined_report_text for keyword in ["RR间期", "瞬时心率", "SDNN", "RMSSD", "HRV技术可用性", "不解释疾病风险"])
            and all(keyword in combined_report_text for keyword in ["不能外推到ST段分析", "不能外推到", "临床决策"]),
            "Report includes RR/HR/SDNN/RMSSD methodology and keeps non-diagnostic boundary",
        )
    )
    checks.append(
        check(
            "Appendix engineering package coverage",
            all(label in combined_report_text for label in ["附录A", "附录B", "附录C", "附录D"])
            and all(keyword in combined_report_text for keyword in ["核心参数", "关键脚本", "结果文件索引", "GitHub仓库", "交互展示页"]),
            "Appendix A-D include parameters, scripts, result index, GitHub and Pages links",
        )
    )
    checks.append(
        check(
            "GitHub and interactive showcase links in report",
            "https://github.com/Lebron-shun/ecg-sampling-adc-noise-robustness" in combined_report_text
            and "https://lebron-shun.github.io/ecg-sampling-adc-noise-robustness/" in combined_report_text,
            "Repository and GitHub Pages URLs found in generated reports",
        )
    )
    required_scripts = [
        "download_data.py",
        "validate_data.py",
        "run_experiments.py",
        "analyze_results.py",
        "monitoring_analysis.py",
        "build_web_data.py",
        "build_report.py",
        "build_pdf.py",
        "audit_project.py",
        "project_core.py",
    ]
    checks.append(
        check(
            "Reproducible source code",
            all((PROJECT_ROOT / "src" / script).exists() for script in required_scripts)
            and (PROJECT_ROOT / "config.json").exists()
            and (PROJECT_ROOT / "requirements.txt").exists(),
            f"{len(required_scripts)} required scripts, config.json, and requirements.txt present",
        )
    )
    checks.append(
        check(
            "Course-guide report coverage",
            all(
                keyword in combined_report_text
                for keyword in ["医学应用背景", "系统设计原理", "方法与实现过程", "结果展示与性能评价", "讨论与改进分析", "参考文献", "附录A", "附录D"]
            ),
            "All required course-report sections and appendices found in generated report text",
        )
    )
    readme_text = (PROJECT_ROOT / "README.md").read_text(encoding="utf-8")
    web_text = (PROJECT_ROOT / "web" / "index.html").read_text(encoding="utf-8")
    checks.append(
        check(
            "README monitoring and showcase entry",
            all(keyword in readme_text for keyword in ["RR/HRV", "GitHub Pages", "FINAL_AUDIT.md", "report/final_report.pdf", "monitoring_analysis.py"])
            and "https://lebron-shun.github.io/ecg-sampling-adc-noise-robustness/" in readme_text,
            "README highlights monitoring extension, report, audit, and GitHub Pages entry",
        )
    )
    web_js = (PROJECT_ROOT / "web" / "app.js").read_text(encoding="utf-8")
    web_data = (PROJECT_ROOT / "web" / "data.js").read_text(encoding="utf-8")
    checks.append(
        check(
            "Interactive web project showcase and figure performance",
            all(keyword in web_text for keyword in ["Project Overview", "Monitoring Scenarios", "医学监护场景", "PDF报告", "项目README", "GitHub仓库"])
            and all(keyword in web_text for keyword in ["可穿戴ECG长时监护", "采样率 × ADC位数", "MIT-BIH与NSTDB", "实验图集"])
            and all(keyword in web_js for keyword in ["preloadFigures", "decode", "figureCache", "selectFigure", "renderScenarios", "renderMonitoring"])
            and all(keyword in web_data for keyword in ["monitoring", "scenarios", "hr_rr_monitoring", "hrv_trend"])
            and all(keyword not in web_text for keyword in ["Course Rubric", "课程评分对照", "5分", "6分", "17/17"]),
            "web/index.html presents monitoring scenarios, not scoring; figure gallery preloads decoded PNG images",
        )
    )

    passed = all(item["passed"] for item in checks)
    audit = {
        "passed": passed,
        "checks_passed": sum(item["passed"] for item in checks),
        "checks_total": len(checks),
        "recommended_configuration": {
            "target_fs_hz": int(recommended["target_fs_hz"]),
            "bits": int(recommended["bits"]),
            "raw_bitrate_bps": int(recommended["raw_bitrate_bps"]),
            "bitrate_reduction_pct": recommended["bitrate_reduction_pct"],
            "clean_pooled_f1_pct": recommended["pooled_f1_pct"],
            "worst_noise_f1_drop_pp_at_snr_ge_6": recommended["worst_noise_f1_drop_pp"],
        },
        "checks": checks,
    }
    AUDIT_JSON.write_text(json.dumps(audit, indent=2, ensure_ascii=False), encoding="utf-8")
    lines = [
        "# Final Project Completion Audit",
        "",
        f"**Overall status:** {'PASS' if passed else 'FAIL'}  ",
        f"**Checks:** {audit['checks_passed']}/{audit['checks_total']} passed  ",
        f"**Recommended configuration:** {int(recommended['target_fs_hz'])} Hz / {int(recommended['bits'])} bit",
        "",
        "| Requirement | Status | Evidence |",
        "|---|---|---|",
    ]
    lines.extend(
        f"| {item['requirement']} | {'PASS' if item['passed'] else 'FAIL'} | {item['evidence'].replace('|', '/')} |"
        for item in checks
    )
    AUDIT_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")
    print(json.dumps(audit, indent=2, ensure_ascii=False))
    if not passed:
        raise RuntimeError("Final project audit failed")


if __name__ == "__main__":
    main()
